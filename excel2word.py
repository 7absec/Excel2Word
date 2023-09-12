"""
NOTE
Program to generate word documents from an excel file.
A python program to generate a word document from a given excel workbook based on predefined columns.
The input excel file should have the same columns given in the example excel file.

### Example Excel File Columns
The input Excel file should have the following columns, just like the example Excel file provided:

|   column1   |   column2   |   column3   |   column4   |   column5   |   column6   |   column7   |   column8   |
|-------------|-------------|-------------|-------------|-------------|-------------|-------------|-------------|
|   Target    | Vulnerability Name |   Severity  |    CVSS     |  Parameter  | Description |   Impact    | Remediation |


### Directory structure

Root Dir (e.g., Eample)
│
├── 1_subdirectory (e.g., eample.com)
│ ├── 1_secondsubdir (e.g., RCE)
│ │ ├── 1_anything.png
│ │ ├── 2_anything.png
│ │ └── ...
│ ├── 2_secondsubdir (e.g., XSS)
│ │ ├── 1_anything.png
│ │ ├── 2_anything.png
│ │ └── ...
│ └── ...
│
├── 2_subdirectory (e.g., google.com)
│ ├── ...
│
└── ...

The output excel file of the prgoram will be same as the example word file.

This program requires python 3.10 or later.
"""

Author = "7absec"

################################################################
#Import the
import pandas as pd
import tkinter as tk
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from tkinter import filedialog,ttk
import openpyxl,os,datetime
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL

################################################################
'''Input section'''

def is_valid_excel_file(file_path):
    if not file_path:
        return False
    _, ext = os.path.splitext(file_path)
    return ext.lower() in ('.xlsx', '.xls', '.xlsm')

def is_valid_directory(directory_path):
    if not directory_path:
        return False
    return os.path.isdir(directory_path)

def get_input_file():
    input_file_entry.delete(0, tk.END)
    file_path = filedialog.askopenfilename(title="Select Excel File", filetypes=[("Excel Files", "*.xlsx *.xls *.xlsm")])
    if is_valid_excel_file(file_path):
        input_file_entry.insert(0, file_path)
    else:
        input_file_entry.delete(0, tk.END)
        input_file_entry.insert(0, "Invalid File")

def get_img_dir():
    img_dir_entry.delete(0, tk.END)
    directory_path = filedialog.askdirectory(title="Select Image Folder")
    if is_valid_directory(directory_path):
        img_dir_entry.insert(0, directory_path)
    else:
        img_dir_entry.delete(0, tk.END)
        img_dir_entry.insert(0, "Invalid Folder")

def get_output_dir():
    output_dir_entry.delete(0, tk.END)
    directory_path = filedialog.askdirectory(title="Select Output Folder")
    if is_valid_directory(directory_path):
        output_dir_entry.insert(0, directory_path)
    else:
        output_dir_entry.delete(0, tk.END)
        output_dir_entry.insert(0, "Invalid Folder")

def submit_entries():
    input_file = input_file_entry.get()
    img_dir = img_dir_entry.get()
    output_dir = output_dir_entry.get()
    client_name = string_input_entry.get()

    # Check if the input file is valid
    if not is_valid_excel_file(input_file):
        input_file_entry.delete(0, tk.END)
        input_file_entry.insert(0, "Invalid File")
        return None

    # Check if the image directory is valid
    if not is_valid_directory(img_dir):
        img_dir_entry.delete(0, tk.END)
        img_dir_entry.insert(0, "Invalid Folder")
        return None

    # Check if the output directory is valid
    if not is_valid_directory(output_dir):
        output_dir_entry.delete(0, tk.END)
        output_dir_entry.insert(0, "Invalid Folder")
        return None

    # Check if the client name is valid (you can add your validation logic here)
    if not client_name:
        string_input_entry.delete(0, tk.END)
        string_input_entry.insert(0, "Invalid Client Name")
        return None

    # Create a dictionary to store the selected values
    selected_values = {
        1: input_file,
        2: img_dir,
        3: output_dir,
        4: client_name
    }
    root.destroy()
    return selected_values

root = tk.Tk()
root.title(f"Excel to Word converter, Author - {Author}")

# Configure the root window
root.geometry("690x350")
root.resizable(False, False)

# Define a custom font
custom_font = ("Arial", 12)

# Create a styled frame for the inputs
frame = ttk.Frame(root, padding=20)
frame.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")

# Add a title label
title_label = ttk.Label(frame, text="Data Input", font=("Arial", 16))
title_label.grid(row=0, column=0, columnspan=2, pady=10)

# Create input widgets with labels
input_file_label = ttk.Label(frame, text="Excel File:", font=custom_font)
input_file_label.grid(row=1, column=0, padx=5, pady=5, sticky="w")
input_file_entry = ttk.Entry(frame, width=40, font=custom_font)
input_file_entry.grid(row=1, column=1, padx=5, pady=5, sticky="w")
input_file_button = ttk.Button(frame, text="Select", command=get_input_file)
input_file_button.grid(row=1, column=2, padx=5, pady=5, sticky="w")

img_dir_label = ttk.Label(frame, text="Image Folder:", font=custom_font)
img_dir_label.grid(row=2, column=0, padx=5, pady=5, sticky="w")
img_dir_entry = ttk.Entry(frame, width=40, font=custom_font)
img_dir_entry.grid(row=2, column=1, padx=5, pady=5, sticky="w")
img_dir_button = ttk.Button(frame, text="Select", command=get_img_dir)
img_dir_button.grid(row=2, column=2, padx=5, pady=5, sticky="w")

output_dir_label = ttk.Label(frame, text="Output Folder:", font=custom_font)
output_dir_label.grid(row=3, column=0, padx=5, pady=5, sticky="w")
output_dir_entry = ttk.Entry(frame, width=40, font=custom_font)
output_dir_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")
output_dir_button = ttk.Button(frame, text="Select", command=get_output_dir)
output_dir_button.grid(row=3, column=2, padx=5, pady=5, sticky="w")

string_input_label = ttk.Label(frame, text="Client Name:", font=custom_font)
string_input_label.grid(row=4, column=0, padx=5, pady=5, sticky="w")
string_input_entry = ttk.Entry(frame, width=40, font=custom_font)
string_input_entry.grid(row=4, column=1, padx=5, pady=5, sticky="w")

# Add a Submit button
submit_button = ttk.Button(frame, text="Submit", command=root.quit)
submit_button.grid(row=5, column=0, columnspan=3, pady=20)

# Center the window on the screen
root.eval('tk::PlaceWindow . center')
root.mainloop()

# After the GUI is closed (when the user clicks "Submit"), you can get the selected values
selected_values = submit_entries()
in_file = selected_values[1]
img_folder = selected_values[2]
out_folder = selected_values[3]
client_name = selected_values[4]


#Set Table header color
def set_cell_background_color(cell, severity):
    if severity == "Critical":
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="9A2008"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)
    elif severity == "High":
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)  # FF0000
    elif severity == "Medium":
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)   # FFC000
    elif severity == "Low":
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="3A7C28"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)   # 3A7C28
    elif severity == "Informational" or severity == "Info":
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="0070C0"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_1)   # 0070C0



#Vulnerability name and severity
def heading(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(16)
            run.font.name = 'Calibri (Headings)'
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
            run.italic = True

################################################################
#Image folder walk function to search for images and add to the appropriate table number
#Insert images

# Function to extract serial numbers from directory names
def extract_serial_number(dir_name):
    try:
        serial_number = int(dir_name.split("_")[0])
        return serial_number
    except ValueError:
        return None

# Function to list all files in a directory and its subdirectories
def list_files_recursively(directory):
    file_list = []
    for root, _, files in os.walk(directory):
        for file in files:
            file_list.append(os.path.join(root, file))
    return file_list

# Function to check for matching worksheet number and TableNum and create a list of image paths
def mulImgCheck(worksheet, TableNum):
    root_dir = img_folder
    image_paths = []
    # Iterate through the subdirectories
    for subdir in os.listdir(root_dir):
        subdir_path = os.path.join(root_dir, subdir)
        # Check if it's a directory
        if os.path.isdir(subdir_path):
            serial_number = extract_serial_number(subdir)
            # Check if the serial number matches the worksheet number
            if serial_number == worksheet:
                for ssubdir in os.listdir(subdir_path):
                    ssdir_path = os.path.join(subdir_path, ssubdir)
                    if os.path.isdir(ssdir_path):
                        dirnum = extract_serial_number(ssubdir)
                        # Check if the TableNum matches any of the subdirectory names
                        if dirnum == TableNum:
                            # List all files in the matched subdirectory and its subdirectories
                            table_dir = root_dir +"\\" + subdir + "\\" + ssubdir
                            files = list_files_recursively(table_dir)
                            # Append the image paths to the list
                            image_paths.extend(files)
    return image_paths


# Function to check for matching worksheet number and TableNum and create a list of image paths
def singleImgCheck(worksheet):
    root_dir = img_folder
    image_paths = []
    # Iterate through the subdirectories
    for subdir in os.listdir(root_dir):
        subdir_path = os.path.join(root_dir, subdir)
        # Check if it's a directory
        if os.path.isdir(subdir_path):
            serial_number = extract_serial_number(subdir)
            # Check if the serial number matches the worksheet number
            if serial_number == worksheet:
                # List all files in the matched subdirectory and its subdirectories
                table_dir = root_dir +"\\" + subdir
                files = list_files_recursively(table_dir)
                # Append the image paths to the list
                image_paths.extend(files)
    return image_paths



#Other parameters
def cell_text_color(cell, severity):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Calibri (Headings)'
            run.bold = True
            if severity == "Critical":
                run.font.color.rgb = RGBColor(154, 32, 8)
            elif severity == "High":
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif severity == "Medium":
                run.font.color.rgb = RGBColor(255, 192, 0)
            elif severity == "Low":
                run.font.color.rgb = RGBColor(58,124,40)
            elif severity == "Informational" or severity == "Info":
                run.font.color.rgb = RGBColor(0,112,192)

#Set cell font size
def cell_font_size(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
            run.font.name = 'Calibri (Headings)'


###############################################################
# Create a new DOCX document
document = Document()

#Load workbook
workbook = openpyxl.load_workbook(in_file)
df = pd.read_excel(in_file)
sheet_names = workbook.sheetnames
def multisheet():
    length = 0
    while length < len(sheet_names):
        sheet = sheet_names[length]
        worksheet = workbook[sheet_names[length]]

        ################################
        #Vulnerability Count section start
        df = pd.read_excel(in_file, sheet_name=sheet)
        unique_values_column1 = df['Target'].unique()
        Target_count = df['Target'].count()

        # Define a custom sorting order
        custom_sort_order = {
            "Critical": 1,
            "High": 2,
            "Medium": 3,
            "Low": 4,
            "Informational": 5,
        }

        # Create a table with headers for the unique values in column 4
        severity_name = list(custom_sort_order.keys())
        table = document.add_table(rows=1, cols=len(severity_name) + 1)
        table.autofit = True

        # Set the header row
        header_row = table.rows[0].cells
        header_row[0].text = "Target"
        shading = parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w')))
        header_row[0]._tc.get_or_add_tcPr().append(shading)
        header_row[0].paragraphs[0].runs[0].bold = True
        header_row[0].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
        header_row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for i, severity in enumerate(severity_name, start=1):
            header_row[i].text = severity
            if severity == "Critical":
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="9A2008"/>'.format(nsdecls('w')))
                header_row[i]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif severity == "High":
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FF0000"/>'.format(nsdecls('w')))
                header_row[i]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif severity == "Medium":
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="FFC000"/>'.format(nsdecls('w')))
                header_row[i]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif severity == "Low":
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="3A7C28"/>'.format(nsdecls('w')))
                header_row[i]._tc.get_or_add_tcPr().append(shading_elm_1)
            elif severity == "Informational":
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="0070C0"/>'.format(nsdecls('w')))
                header_row[i]._tc.get_or_add_tcPr().append(shading_elm_1)

            header_row[i].paragraphs[0].runs[0].bold = True
            header_row[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(200,200,200)
            header_row[i].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            header_row[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for value in unique_values_column1:
            # Filter the DataFrame for the current unique value in column 1
            filtered_df = df[df['Target'] == value]
            # Count the occurrences of each value in column 4 for the filtered DataFrame
            counts = filtered_df['Severity'].value_counts()
            # Add the Target and count rows
            Target_row = table.add_row().cells
            Target_row[0].text = value
            for i, severity in enumerate(severity_name, start=1):
                Target_row[i].text = str(counts.get(severity, 0))

        table.style = "Table Grid"
        empty_paragraph = document.add_paragraph()
        empty_paragraph.space_after = Pt(12)

        ################################################################
        #Vulnerability heading table section

        headers = worksheet[1]
        header_values = [cell.value for cell in headers]
        header_to_include = ['Vulnerability Name', 'Severity','CVSS']
        num_cols = len(header_to_include) + 1
        table = document.add_table(rows=1, cols = num_cols) #Initialize an empty table
        table.autofit = True

        #Set the tbale header row
        header_row = table.rows[0].cells
        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w')))
        header_row[0].text = 'Sr No.'
        header_row[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        header_row[0].paragraphs[0].runs[0].bold = True
        header_row[0].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
        header_row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for i, header_text in enumerate(header_values):
            if header_text in header_to_include:
                index_in_table = header_to_include.index(header_text) + 1
                header_cell = header_row[index_in_table]
                header_cell.text = header_text

                #Apply stying to the header row
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="000080"/>'.format(nsdecls('w')))
                header_cell._tc.get_or_add_tcPr().append(shading_elm_1)
                header_cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for paragraph in header_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.bold = True


        #Iterate through the rows in the Ecxcel sheet
        serial_number = 1
        for row in worksheet.iter_rows(min_row=2, values_only=True):
            row_cells = table.add_row().cells
            row_cells[0].text = str(serial_number)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            #Populate the specified columns
            for i, header_text in enumerate(header_values):
                if header_text in header_to_include:
                    if header_text == "CVSS":
                        index_in_table = header_to_include.index(header_text) + 1
                        # cvss = row[i].split("-")
                        cvss = row[i]
                        # score = cvss[1:]
                        row_cells[index_in_table].text = str(cvss[:3])
                        row_cells[2].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        row_cells[3].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                        row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    else:
                        index_in_table = header_to_include.index(header_text) + 1
                        row_cells[index_in_table].text = str(row[i])
                        row_cells[2].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                        row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        row_cells[3].paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
                        row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                if header_text == 'Severity':
                    set_cell_background_color(row_cells[index_in_table], row[i])
            serial_number += 1
            table.style = "Table Grid"

        #Bold the Serial Number field
        i = 0
        for cell in table.columns[0].cells:
                cell.width = Pt(50)
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(12)

        #Vulnerability name cell
        for cell in table.columns[1].cells:
            cell.width = Pt(200)
            cell.paragraphs[0].runs[0].font.size = Pt(12)

        #Change CVSS color to gray and bold the font
        for cell in table.columns[2].cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(200, 200, 200)

        #CVSS cell
        for cell in table.columns[3].cells:
            cell.paragraphs[0].runs[0].font.size = Pt(12)

        document.add_page_break()


        ################################################################
        #Vulnerability Table section

        # Iterate through the rows in the Excel sheet (excluding the first row)
        header = ['Target','Vulnerability Name','Severity','CVSS','Organization','Parameter','Description','Impact','Remediation','PoC']


        l=1
        while l < Target_count+1 :
            k = []
            i = 0
            for i in range(1,10):
                j = worksheet.cell(l+1,i)
                k.append(j.value)

            table = document.add_table(8,2)
            table.style = "Table Grid"
            table.autofit = True
            for cell in table.columns[0].cells:
                cell.width = Pt(50)
            for row in table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.border_width = Pt(1)
                    cell.border_color = (0,0,0)
                    cell.border_top = cell.border_bottom = cell.border_left = cell.border_right = True

            #Add values to the table
            c00 = table.cell(0,0)
            c00.text = str(l)
            c00.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading(c00)
            c00.paragraphs[0].runs[0].font.size = Pt(24)

            #Vulnerability Name
            c01 = table.cell(0,1)
            c01.text = "\nVulnerability Name: " + str(k[1]) + "\nSeverity: " + str(k[2]) + "\n"
            heading(c01)
            if str(k[2]):
                set_cell_background_color(c00, k[2])
                set_cell_background_color(c01, k[2])

            #CVSS value
            c10 = table.cell(1,0)
            c10.text = header[3]
            cell_text_color(c10, k[2])
            c11 = table.cell(1,1)
            c11.text = str(k[3])
            cell_font_size(c11)

            #Organization Name
            c20 = table.cell(2,0)
            c20.text = header[4]
            cell_text_color(c20, k[2])
            c21 = table.cell(2,1)
            c21.text = "client_name"
            cell_font_size(c21)

            #Vulnerable parameter
            c30 = table.cell(3,0)
            c30.text = header[5]
            cell_text_color(c30, k[2])
            c31 = table.cell(3,1)
            c31.text = str(k[4])
            cell_font_size(c31)

            #DescrTargettion
            c40 = table.cell(4,0)
            c40.text = header[6]
            cell_text_color(c40, k[2])
            c41 = table.cell(4,1)
            c41.text = str(k[5]) + "\n"
            cell_font_size(c41)

            #Impact
            c50 = table.cell(5,0)
            c50.text = header[7]
            cell_text_color(c50, k[2])
            c51 = table.cell(5,1)
            c51.text = str(k[6]) + "\n"
            cell_font_size(c51)

            #Remediation
            c60 = table.cell(6,0)
            c60.text = header[8]
            cell_text_color(c60, k[2])
            c61 = table.cell(6,1)
            c61.text = str(k[7]) + "\n"
            cell_font_size(c61)

            #PoC
            c70 = table.cell(7,0)
            c70.text = header[9]
            cell_text_color(c70, k[2])
            c71 = table.cell(7,1)
            c71.text = "Step 1:"
            cell_font_size(c71)

            #Check the number of sheets
            if len(sheet_names) == 1:
                imgs = singleImgCheck(l)

            else:
                imgs = mulImgCheck(length+1, l)

            paranum = 0
            for img in imgs:
                #Add steps
                if paranum:
                    paragraph = c71.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    new_text = "\nStep" + str(paranum+1) + ":"
                    paragraph.add_run(new_text)
                    cell_font_size(c71)

                #insert image
                img_path = img
                cell_paragraph = c71.paragraphs[paranum]
                run = cell_paragraph.add_run()
                run.add_picture(img_path, width=Inches(4.8), height=Inches(2.5))
                paranum += 1

            #Break the page after every table
            document.add_page_break()
            k.clear()  #Clear the list values
            l+=1
        length += 1

multisheet()


# Save the DOCX document
timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H%M%S")
doc_file = out_folder + "//" + f"Output_{timestamp}.docx"
document.save(doc_file)

workbook.close()
print("Operation completed successfully")